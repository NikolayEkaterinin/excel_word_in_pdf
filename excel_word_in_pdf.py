import pandas as pd
import matplotlib.pyplot as plt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from docx import Document
import io
import os

PAGE_WIDTH, PAGE_HEIGHT = A4
ROWS_PER_PAGE = 40

def read_excel_file(file_path: str) -> pd.DataFrame:
    """Читаем Excel и возвращаем DataFrame"""
    return pd.read_excel(file_path)

def read_word_file(file_path: str) -> pd.DataFrame:
    """Читаем Word (.docx), вытаскиваем таблицы и параграфы"""
    doc = Document(file_path)
    all_parts = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if len(rows) > 1:
            df = pd.DataFrame(rows[1:], columns=rows[0])
        else:
            df = pd.DataFrame(rows)
        all_parts.append(df)

    text_data = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if text_data:
        df_text = pd.DataFrame(text_data, columns=["Text"])
        all_parts.append(df_text)

    if all_parts:
        return pd.concat(all_parts, ignore_index=True, sort=False)
    else:
        raise ValueError("В Word-документе не найдено текста или таблиц")

def build_pdf(df: pd.DataFrame, pdf_file: str):
    """Создание PDF из DataFrame с водяным знаком"""
    c = canvas.Canvas(pdf_file, pagesize=A4)

    for start in range(0, len(df), ROWS_PER_PAGE):
        end = start + ROWS_PER_PAGE
        chunk = df.iloc[start:end]

        fig, ax = plt.subplots(figsize=(12, len(chunk) * 0.3 + 1))
        ax.axis("off")

        table = ax.table(
            cellText=chunk.values,
            colLabels=chunk.columns,
            cellLoc="center",
            loc="center"
        )

        table.auto_set_font_size(False)
        table.set_fontsize(9)
        table.scale(1, 1.2)

        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)

        image = ImageReader(buf)
        c.drawImage(image, 30, 100, width=PAGE_WIDTH - 60,
                    preserveAspectRatio=True, mask="auto")

        # водяной знак
        c.saveState()
        c.setFont("Helvetica-Bold", 50)
        try:
            c.setFillColorRGB(0.9, 0.9, 0.9, alpha=0.3)
        except TypeError:
            c.setFillGray(0.9)
        c.translate(PAGE_WIDTH / 2, PAGE_HEIGHT / 2)
        c.rotate(45)
        c.drawCentredString(0, 0, "M&N Digital")
        c.restoreState()

        c.showPage()

    c.save()

if __name__ == "__main__":
    input_file = "input.docx"
    pdf_file = "demo_with_watermark.pdf"

    ext = os.path.splitext(input_file)[1].lower()
    if ext in [".xlsx", ".xls"]:
        df = read_excel_file(input_file)
    elif ext == ".docx":
        df = read_word_file(input_file)
    else:
        raise ValueError("Поддерживаются только .xlsx, .xls и .docx")

    build_pdf(df, pdf_file)
    print(f"PDF с водяным знаком сохранён как {pdf_file}")
